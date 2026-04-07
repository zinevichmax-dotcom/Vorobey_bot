"""
Сборщик DOCX-отчёта по результатам сравнения.
Выделяет: зелёным — добавленное, красным — удалённое,
жёлтым — изменённое.
"""

from __future__ import annotations

import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor


# Цвета
COLOR_ADDED = RGBColor(0x00, 0x80, 0x00)  # зелёный
COLOR_REMOVED = RGBColor(0xCC, 0x00, 0x00)  # красный
COLOR_NORMAL = RGBColor(0x33, 0x33, 0x33)  # тёмно-серый
COLOR_LABEL = RGBColor(0x66, 0x66, 0x66)  # серый для меток


def build_diff_report(diff_data: dict, output_path: str, name_a: str = "Файл A", name_b: str = "Файл B") -> str:
    """
    Собирает DOCX-отчёт по diff.
    """
    doc = Document()

    # Стиль по умолчанию
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    # --- Заголовок ---
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("ОТЧЁТ О СРАВНЕНИИ ДОКУМЕНТОВ")
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"

    doc.add_paragraph()

    # --- Статистика ---
    stats = diff_data["stats"]
    info = doc.add_paragraph()
    info.add_run(f"Документ A: {name_a} ({diff_data['file_a_paragraphs']} абзацев)\n").font.size = Pt(11)
    info.add_run(f"Документ B: {name_b} ({diff_data['file_b_paragraphs']} абзацев)\n").font.size = Pt(11)
    info.add_run(f"\nИзменено: {stats['modified']}  |  ").font.size = Pt(11)
    info.add_run(f"Добавлено: {stats['added']}  |  ").font.size = Pt(11)
    info.add_run(f"Удалено: {stats['removed']}  |  ").font.size = Pt(11)
    info.add_run(f"Без изменений: {stats['unchanged']}").font.size = Pt(11)

    doc.add_paragraph()
    _add_separator(doc)

    # --- Детали изменений ---
    change_num = 0
    for c in diff_data["changes"]:
        ctype = c["type"]

        if ctype == "unchanged":
            continue

        change_num += 1

        # Заголовок изменения
        header = doc.add_paragraph()
        label_run = header.add_run(f"Изменение #{change_num}  ")
        label_run.bold = True
        label_run.font.size = Pt(11)
        label_run.font.name = "Times New Roman"

        type_text = {"added": "[ДОБАВЛЕНО]", "removed": "[УДАЛЕНО]", "modified": "[ИЗМЕНЕНО]"}
        type_color = {"added": COLOR_ADDED, "removed": COLOR_REMOVED, "modified": RGBColor(0xCC, 0x88, 0x00)}

        type_run = header.add_run(type_text.get(ctype, ctype))
        type_run.bold = True
        type_run.font.color.rgb = type_color.get(ctype, COLOR_NORMAL)
        type_run.font.size = Pt(11)

        # Содержимое
        if ctype == "modified" and c.get("inline_diff"):
            # Показываем детальный inline diff
            _add_label(doc, f"Было (п.{c.get('para_a', '?')}):")
            old_para = doc.add_paragraph()
            old_para.paragraph_format.left_indent = Pt(20)
            for part in c["inline_diff"]:
                run = old_para.add_run(part["text"] + " ")
                run.font.size = Pt(11)
                run.font.name = "Times New Roman"
                if part["type"] == "removed":
                    run.font.color.rgb = COLOR_REMOVED
                    run.font.strike = True
                elif part["type"] == "added":
                    run.font.color.rgb = COLOR_ADDED
                    run.bold = True
                else:
                    run.font.color.rgb = COLOR_NORMAL

            _add_label(doc, f"Стало (п.{c.get('para_b', '?')}):")
            new_para = doc.add_paragraph()
            new_para.paragraph_format.left_indent = Pt(20)
            for part in c["inline_diff"]:
                if part["type"] != "removed":
                    run = new_para.add_run(part["text"] + " ")
                    run.font.size = Pt(11)
                    run.font.name = "Times New Roman"
                    if part["type"] == "added":
                        run.font.color.rgb = COLOR_ADDED
                        run.bold = True
                    else:
                        run.font.color.rgb = COLOR_NORMAL

        elif ctype == "removed":
            _add_label(doc, f"Удалено из п.{c.get('para_a', '?')}:")
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(20)
            run = p.add_run(c["text_a"])
            run.font.color.rgb = COLOR_REMOVED
            run.font.strike = True
            run.font.size = Pt(11)

        elif ctype == "added":
            _add_label(doc, f"Добавлено в п.{c.get('para_b', '?')}:")
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(20)
            run = p.add_run(c["text_b"])
            run.font.color.rgb = COLOR_ADDED
            run.font.size = Pt(11)

        _add_separator(doc)

    # Сохраняем
    os.makedirs(os.path.dirname(output_path) if os.path.dirname(output_path) else ".", exist_ok=True)
    doc.save(output_path)
    return output_path


def _add_label(doc, text: str):
    """Добавляет серую метку."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.color.rgb = COLOR_LABEL
    run.font.size = Pt(10)
    run.italic = True


def _add_separator(doc):
    """Добавляет визуальный разделитель."""
    p = doc.add_paragraph()
    run = p.add_run("—" * 60)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    run.font.size = Pt(8)


if __name__ == "__main__":
    import json
    import sys

    if len(sys.argv) < 2:
        print("Использование: python diff_report_builder.py diff_result.json [output.docx]")
        raise SystemExit(1)

    with open(sys.argv[1]) as f:
        diff_data = json.load(f)

    output = sys.argv[2] if len(sys.argv) > 2 else "diff_report.docx"
    build_diff_report(diff_data, output)
    print(f"Сохранено: {output}")

