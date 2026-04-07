"""
Сборщик DOCX для допсоглашения.
Берёт структуру от Claude и собирает юридический документ.
"""

from __future__ import annotations

import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def build_agreement_docx(agreement_data: dict, output_path: str) -> str:
    """
    Собирает DOCX допсоглашения.

    agreement_data = {
        "title": "ДОПОЛНИТЕЛЬНОЕ СОГЛАШЕНИЕ...",
        "preamble": "...",
        "clauses": [{"num": 1, "text": "..."}, ...],
        "closing": "...",
    }
    """
    doc = Document()

    # --- Стили ---
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)
    paragraph_format = style.paragraph_format
    paragraph_format.space_after = Pt(0)
    paragraph_format.space_before = Pt(0)
    paragraph_format.line_spacing = 1.5

    # --- Заголовок ---
    for line in (agreement_data.get("title") or "").split("\n"):
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run(line.strip())
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(14) if "ДОПОЛНИТЕЛЬНОЕ" in line else Pt(12)

    doc.add_paragraph()  # пустая строка

    # --- Преамбула ---
    preamble_para = doc.add_paragraph()
    preamble_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = preamble_para.add_run(agreement_data.get("preamble") or "")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    doc.add_paragraph()  # пустая строка

    # --- Пункты ---
    for clause in agreement_data.get("clauses") or []:
        clause_para = doc.add_paragraph()
        clause_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        num_run = clause_para.add_run(f"{clause['num']}. ")
        num_run.bold = True
        num_run.font.name = "Times New Roman"
        num_run.font.size = Pt(12)

        text_run = clause_para.add_run(clause["text"])
        text_run.font.name = "Times New Roman"
        text_run.font.size = Pt(12)

    doc.add_paragraph()  # пустая строка

    # --- Заключительные положения ---
    for line in (agreement_data.get("closing") or "").split("\n"):
        if line.strip():
            closing_para = doc.add_paragraph()
            closing_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = closing_para.add_run(line)
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)

    # --- Сохраняем ---
    os.makedirs(os.path.dirname(output_path) if os.path.dirname(output_path) else ".", exist_ok=True)
    doc.save(output_path)
    return output_path

