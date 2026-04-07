"""
Сборщик DOCX-справки по результатам compliance проверки.
"""

from __future__ import annotations

import os
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor


COLOR_APPROVED = RGBColor(0x00, 0x80, 0x00)
COLOR_REJECTED = RGBColor(0xCC, 0x00, 0x00)
COLOR_WARNING = RGBColor(0xCC, 0x88, 0x00)
COLOR_INFO = RGBColor(0x33, 0x66, 0x99)
COLOR_NORMAL = RGBColor(0x33, 0x33, 0x33)


def build_compliance_report(result: dict, output_path: str, document_name: str = "документ") -> str:
    """Собирает DOCX-справку."""
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    # --- Заголовок ---
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("СПРАВКА О ПРОВЕРКЕ ДОКУМЕНТА")
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"

    # --- Дата ---
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_para.add_run(datetime.now().strftime("%d.%m.%Y")).font.size = Pt(11)

    doc.add_paragraph()

    # --- Вердикт ---
    verdict_para = doc.add_paragraph()
    verdict_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    verdict_text = result.get("verdict", "НЕ ОПРЕДЕЛЕНО")
    approved = result.get("approved", False)

    run = verdict_para.add_run(verdict_text)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = COLOR_APPROVED if approved else COLOR_REJECTED

    doc.add_paragraph()

    # --- Резюме ---
    summary = result.get("summary", "")
    if summary:
        p = doc.add_paragraph()
        run = p.add_run(f"Проверяемый документ: {document_name}")
        run.bold = True
        run.font.size = Pt(12)

        p = doc.add_paragraph()
        p.add_run(summary).font.size = Pt(12)

    doc.add_paragraph()

    # --- Нарушения ---
    violations = result.get("violations", [])
    if violations:
        header = doc.add_paragraph()
        run = header.add_run(f"ВЫЯВЛЕННЫЕ НАРУШЕНИЯ ({len(violations)})")
        run.bold = True
        run.font.size = Pt(13)

        doc.add_paragraph()

        for i, v in enumerate(violations, 1):
            severity = v.get("severity", "info")
            severity_color = {
                "critical": COLOR_REJECTED,
                "warning": COLOR_WARNING,
                "info": COLOR_INFO,
            }.get(severity, COLOR_NORMAL)

            severity_label = {
                "critical": "КРИТИЧЕСКОЕ",
                "warning": "ПРЕДУПРЕЖДЕНИЕ",
                "info": "ИНФОРМАЦИЯ",
            }.get(severity, str(severity).upper())

            # Заголовок нарушения
            p = doc.add_paragraph()
            num_run = p.add_run(f"#{i}  ")
            num_run.bold = True

            sev_run = p.add_run(f"[{severity_label}]")
            sev_run.bold = True
            sev_run.font.color.rgb = severity_color

            # Пункт документа
            if v.get("document_clause"):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(1)
                p.add_run("Пункт документа: ").bold = True
                p.add_run(v["document_clause"])

            # Ссылка на норму
            if v.get("regulatory_reference"):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(1)
                p.add_run("Нормативная основа: ").bold = True
                p.add_run(v["regulatory_reference"])

            # Описание
            if v.get("description"):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(1)
                p.add_run("Суть нарушения: ").bold = True
                p.add_run(v["description"])

            # Рекомендация
            if v.get("recommendation"):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(1)
                rec_run = p.add_run("Рекомендация: ")
                rec_run.bold = True
                rec_run.font.color.rgb = COLOR_INFO
                p.add_run(v["recommendation"])

            # Разделитель
            sep = doc.add_paragraph()
            sep.add_run("—" * 50).font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

    else:
        p = doc.add_paragraph()
        run = p.add_run("Нарушений не выявлено.")
        run.font.color.rgb = COLOR_APPROVED
        run.bold = True

    # --- Примечания ---
    notes = result.get("notes", [])
    if notes:
        doc.add_paragraph()
        header = doc.add_paragraph()
        header.add_run("ПРИМЕЧАНИЯ").bold = True

        for note in notes:
            p = doc.add_paragraph()
            p.add_run(f"• {note}").font.size = Pt(11)

    # --- Дисклеймер ---
    doc.add_paragraph()
    disc = doc.add_paragraph()
    disc_run = disc.add_run(
        "Данная справка подготовлена автоматически с использованием AI-анализа "
        "и носит рекомендательный характер. Окончательное решение об одобрении "
        "документа принимается уполномоченным лицом."
    )
    disc_run.italic = True
    disc_run.font.size = Pt(9)
    disc_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    os.makedirs(os.path.dirname(output_path) if os.path.dirname(output_path) else ".", exist_ok=True)
    doc.save(output_path)
    return output_path

