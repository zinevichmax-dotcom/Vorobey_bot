"""
Хранилище нормативных документов.
Загружаются один раз, извлекается текст, кешируется.
"""

from __future__ import annotations

import hashlib
import json
import os
import shutil

from docx import Document


STORE_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "data", "compliance_docs"))

DOC_TYPE_LABELS: dict[str, str] = {
    "fz_208": "ФЗ-208",
    "fz_14": "ФЗ-14",
    "charter": "Устав",
    "corporate_agreement": "Корпоративный договор",
}

VALID_DOC_TYPES: tuple[str, ...] = ("fz_208", "fz_14", "charter", "corporate_agreement")

# Проходы compliance проверки: какие НПА сравниваем с проверяемым документом
PASS_GROUPS: dict[int, tuple[str, ...]] = {
    1: ("fz_208",),
    2: ("fz_14", "charter"),
    3: ("corporate_agreement",),
}

def _approx_tokens_from_text(paragraphs: list[str]) -> int:
    """
    Грубая оценка токенов: ~4 символа на токен для кириллицы/латиницы.
    Используется только для контроля лимитов контекста.
    """
    char_count = sum(len(p) for p in paragraphs)
    return max(0, int(char_count / 4))


def get_total_tokens() -> dict:
    """
    Считает примерный токен-бюджет по проходам compliance.
    Возвращает структуру:
    {
        "passes": { "1": {"total_tokens": int, "fits_200k": bool}, ... },
        "total": int
    }
    """
    docs = get_regulatory_texts()
    passes: dict[str, dict] = {}
    total = 0

    for pass_num, doc_types in PASS_GROUPS.items():
        pass_tokens = 0
        for doc_type in doc_types:
            data = docs.get(doc_type)
            if not data:
                continue
            pass_tokens += _approx_tokens_from_text(data.get("text", []))
        passes[str(pass_num)] = {
            "total_tokens": pass_tokens,
            "fits_200k": pass_tokens <= 200_000,
        }
        total += pass_tokens

    return {"passes": passes, "total": total}


def init_store():
    """Создаёт папку хранилища если нет."""
    os.makedirs(STORE_DIR, exist_ok=True)
    os.makedirs(os.path.join(STORE_DIR, "originals"), exist_ok=True)
    os.makedirs(os.path.join(STORE_DIR, "extracted"), exist_ok=True)


def upload_regulatory_doc(file_path: str, doc_type: str, doc_name: str) -> dict:
    """
    Загружает нормативный документ в хранилище.
    doc_type: "fz_208" | "fz_14" | "charter" | "corporate_agreement"
    Поддерживает .docx и .odt.
    """
    init_store()

    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".docx":
        text = _extract_text_docx(file_path)
    elif ext == ".odt":
        text = _extract_text_odt(file_path)
    else:
        raise ValueError(f"Неподдерживаемый формат: {ext}")

    # Считаем хеш для дедупликации
    with open(file_path, "rb") as f:
        file_hash = hashlib.md5(f.read()).hexdigest()[:12]

    # Сохраняем метаданные
    meta = {
        "doc_type": doc_type,
        "doc_type_label": DOC_TYPE_LABELS.get(doc_type, doc_type),
        "doc_name": doc_name,
        "file_hash": file_hash,
        "paragraphs": len(text),
        "char_count": sum(len(t) for t in text),
    }
    meta["file_format"] = ext

    # Сохраняем извлечённый текст
    text_path = os.path.join(STORE_DIR, "extracted", f"{doc_type}.json")
    with open(text_path, "w", encoding="utf-8") as f:
        json.dump({"meta": meta, "text": text}, f, ensure_ascii=False, indent=2)

    # Копируем оригинал
    orig_path = os.path.join(STORE_DIR, "originals", f"{doc_type}{ext}")
    shutil.copy2(file_path, orig_path)

    return meta


def get_regulatory_texts() -> dict:
    """
    Возвращает тексты всех загруженных нормативных документов.
    {
        "federal_law": {"meta": {...}, "text": [...]},
        "charter": {"meta": {...}, "text": [...]},
        "corporate_agreement": {"meta": {...}, "text": [...]},
    }
    """
    init_store()
    extracted_dir = os.path.join(STORE_DIR, "extracted")
    docs: dict = {}

    for filename in os.listdir(extracted_dir):
        if filename.endswith(".json"):
            doc_type = filename.replace(".json", "")
            with open(os.path.join(extracted_dir, filename), "r", encoding="utf-8") as f:
                docs[doc_type] = json.load(f)

    return docs


def get_docs_for_pass(pass_num: int) -> dict:
    """
    Возвращает загруженные документы, которые участвуют в указанном проходе.
    {
        "<doc_type>": {"meta": {...}, "text": [...]},
        ...
    }
    """
    docs = get_regulatory_texts()
    allowed = set(PASS_GROUPS.get(pass_num, ()))
    if not allowed:
        return {}
    return {dtype: data for dtype, data in docs.items() if dtype in allowed}


def get_regulatory_summary() -> list:
    """Список загруженных документов с метаданными."""
    docs = get_regulatory_texts()
    return [
        {
            "doc_type": dtype,
            "doc_name": data["meta"]["doc_name"],
            "paragraphs": data["meta"]["paragraphs"],
            "char_count": data["meta"]["char_count"],
        }
        for dtype, data in docs.items()
    ]


def _extract_text_docx(docx_path: str) -> list[str]:
    """Извлекает текст из DOCX по абзацам."""
    doc = Document(docx_path)
    paragraphs: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    # Таблицы тоже
    for table in doc.tables:
        for row in table.rows:
            row_texts: list[str] = []
            for cell in row.cells:
                if cell.text.strip():
                    row_texts.append(cell.text.strip())
            if row_texts:
                paragraphs.append(" | ".join(row_texts))

    return paragraphs


def _extract_text_odt(odt_path: str) -> list[str]:
    """Извлекает текст из ODT (OpenDocument). ODT = ZIP с XML внутри."""
    import zipfile
    import xml.etree.ElementTree as ET

    with zipfile.ZipFile(odt_path, "r") as z:
        with z.open("content.xml") as f:
            tree = ET.parse(f)

    root = tree.getroot()
    texts = []
    for elem in root.iter():
        if elem.text and elem.text.strip():
            texts.append(elem.text.strip())
        if elem.tail and elem.tail.strip():
            texts.append(elem.tail.strip())

    # Объединяем мелкие фрагменты в абзацы
    paragraphs = []
    current = []
    for t in texts:
        current.append(t)
        if len(t) > 80 or t.endswith((".", ":", ";")):
            merged = " ".join(current)
            if len(merged) > 5:
                paragraphs.append(merged)
            current = []
    if current:
        merged = " ".join(current)
        if len(merged) > 5:
            paragraphs.append(merged)

    return paragraphs

