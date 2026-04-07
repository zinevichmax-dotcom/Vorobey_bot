"""
Сравнение двух DOCX файлов.
Извлекает текст → сравнивает по абзацам → возвращает diff.
Использует difflib (stdlib, без зависимостей).
"""

from __future__ import annotations

import difflib
from enum import Enum

from docx import Document


class ChangeType(str, Enum):
    ADDED = "added"
    REMOVED = "removed"
    MODIFIED = "modified"
    UNCHANGED = "unchanged"


def extract_paragraphs(docx_path: str) -> list[dict]:
    """
    Извлекает абзацы из DOCX.
    Возвращает список: [{"index": 0, "text": "...", "style": "Heading 1"}, ...]
    """
    doc = Document(docx_path)
    paragraphs: list[dict] = []

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        paragraphs.append(
            {
                "index": i,
                "text": text,
                "style": para.style.name if para.style else "Normal",
            }
        )

    # Таблицы тоже
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                text = cell.text.strip()
                if text:
                    paragraphs.append(
                        {
                            "index": f"table_{table_idx}_r{row_idx}_c{cell_idx}",
                            "text": text,
                            "style": f"Table[{table_idx}][{row_idx}][{cell_idx}]",
                        }
                    )

    return paragraphs


def compare_documents(path_a: str, path_b: str) -> dict:
    """
    Сравнивает два DOCX файла.
    Возвращает структурированный diff.
    """
    paras_a = extract_paragraphs(path_a)
    paras_b = extract_paragraphs(path_b)

    texts_a = [p["text"] for p in paras_a]
    texts_b = [p["text"] for p in paras_b]

    # Используем SequenceMatcher для умного сравнения
    matcher = difflib.SequenceMatcher(None, texts_a, texts_b)

    changes: list[dict] = []
    stats = {"added": 0, "removed": 0, "modified": 0, "unchanged": 0}

    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "equal":
            for k in range(i1, i2):
                changes.append(
                    {
                        "type": ChangeType.UNCHANGED,
                        "text_a": texts_a[k],
                        "text_b": texts_a[k],
                        "para_a": k + 1,
                        "para_b": j1 + (k - i1) + 1,
                    }
                )
                stats["unchanged"] += 1

        elif tag == "replace":
            # Модификация: текст изменён
            pairs_a = list(range(i1, i2))
            pairs_b = list(range(j1, j2))
            max_len = max(len(pairs_a), len(pairs_b))

            for k in range(max_len):
                text_old = texts_a[pairs_a[k]] if k < len(pairs_a) else None
                text_new = texts_b[pairs_b[k]] if k < len(pairs_b) else None

                if text_old and text_new:
                    # Детальный diff внутри абзаца
                    inline_diff = _inline_diff(text_old, text_new)
                    changes.append(
                        {
                            "type": ChangeType.MODIFIED,
                            "text_a": text_old,
                            "text_b": text_new,
                            "para_a": pairs_a[k] + 1 if k < len(pairs_a) else None,
                            "para_b": pairs_b[k] + 1 if k < len(pairs_b) else None,
                            "inline_diff": inline_diff,
                        }
                    )
                    stats["modified"] += 1
                elif text_old:
                    changes.append(
                        {
                            "type": ChangeType.REMOVED,
                            "text_a": text_old,
                            "text_b": None,
                            "para_a": pairs_a[k] + 1,
                            "para_b": None,
                        }
                    )
                    stats["removed"] += 1
                elif text_new:
                    changes.append(
                        {
                            "type": ChangeType.ADDED,
                            "text_a": None,
                            "text_b": text_new,
                            "para_a": None,
                            "para_b": pairs_b[k] + 1,
                        }
                    )
                    stats["added"] += 1

        elif tag == "insert":
            for k in range(j1, j2):
                changes.append(
                    {
                        "type": ChangeType.ADDED,
                        "text_a": None,
                        "text_b": texts_b[k],
                        "para_a": None,
                        "para_b": k + 1,
                    }
                )
                stats["added"] += 1

        elif tag == "delete":
            for k in range(i1, i2):
                changes.append(
                    {
                        "type": ChangeType.REMOVED,
                        "text_a": texts_a[k],
                        "text_b": None,
                        "para_a": k + 1,
                        "para_b": None,
                    }
                )
                stats["removed"] += 1

    return {
        "file_a_paragraphs": len(paras_a),
        "file_b_paragraphs": len(paras_b),
        "stats": stats,
        "changes": changes,
    }


def _inline_diff(text_a: str, text_b: str) -> list[dict]:
    """
    Детальный diff внутри одного абзаца — слово за словом.
    Возвращает: [{"type": "equal|added|removed", "text": "..."}]
    """
    words_a = text_a.split()
    words_b = text_b.split()

    matcher = difflib.SequenceMatcher(None, words_a, words_b)
    result: list[dict] = []

    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "equal":
            result.append({"type": "equal", "text": " ".join(words_a[i1:i2])})
        elif tag == "replace":
            result.append({"type": "removed", "text": " ".join(words_a[i1:i2])})
            result.append({"type": "added", "text": " ".join(words_b[j1:j2])})
        elif tag == "insert":
            result.append({"type": "added", "text": " ".join(words_b[j1:j2])})
        elif tag == "delete":
            result.append({"type": "removed", "text": " ".join(words_a[i1:i2])})

    return result


if __name__ == "__main__":
    import json
    import sys

    if len(sys.argv) < 3:
        print("Использование: python docx_diff.py файл_A.docx файл_B.docx")
        raise SystemExit(1)

    result = compare_documents(sys.argv[1], sys.argv[2])

    # Показываем только изменения
    print(f"Файл A: {result['file_a_paragraphs']} абзацев")
    print(f"Файл B: {result['file_b_paragraphs']} абзацев")
    print(f"Статистика: {result['stats']}")
    print()

    for c in result["changes"]:
        if c["type"] != ChangeType.UNCHANGED:
            print(f"[{c['type']}]")
            if c.get("text_a"):
                print(f"  A: {c['text_a'][:80]}")
            if c.get("text_b"):
                print(f"  B: {c['text_b'][:80]}")
            print()

